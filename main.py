import os
from dotenv import load_dotenv
load_dotenv()
import sys
import json
import uuid
import re
import base64
from datetime import datetime
from typing import List, Optional, Dict, Any

import httpx
import psycopg2
from psycopg2.extras import RealDictCursor
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import JSONResponse

from gigachat import GigaChat
from gigachat.models import Chat, Messages, MessagesRole

app = FastAPI(title="MedAI Backend GigaChat-2-Max")

# ==============================================================================
# ПЕРЕМЕННЫЕ ОКРУЖЕНИЯ И НАСТРОЙКИ
# ==============================================================================
GIGACHAT_AUTH_KEY = os.getenv("GIGACHAT_AUTH_KEY", "")
GIGACHAT_MODEL = os.getenv("GIGACHAT_MODEL", "GigaChat-2-Max")
GIGACHAT_SCOPE = os.getenv("GIGACHAT_SCOPE", "GIGACHAT_API_PERS")

YANDEX_API_KEY = os.getenv("YANDEX_CLOUD_API_KEY", "")
YANDEX_FOLDER_ID = os.getenv("YANDEX_CLOUD_FOLDER", "")

# РАСШИРЕННЫЙ СПИСОК СТОП-СЛОВ (ВКЛЮЧАЯ МЕСТОИМЕНИЯ И СОЮЗЫ)
STOP_WORDS = {
    'покажи', 'показать', 'объясни', 'расскажи', 'где', 'находится', 'как', 'выглядит',
    'фото', 'картинка', 'картинку', 'картинки', 'фотка', 'фотке', 'фотку', 'фотки', 'фотография', 'фотографии',
    'рисунок', 'рисунки', 'рисунке', 'схема', 'схемы', 'схеме', 'иллюстрация', 'иллюстрации', 'изображение', 'изображения',
    'описание', 'строение', 'функции', 'полное', 'подробное', 'а', 'будет', 'на', 'латинском', 'латыни', 'переведи',
    'причины', 'причина', 'вывиха', 'вывих', 'патология', 'этиология', 'симптом', 'симптомы', 'лечение', 'почему',
    'факторы', 'местонахождение', 'расположение', 'положение', 'разбор', 'мне', 'его', 'её',
    'их', 'у', 'в', 'для', 'о', 'об', 'про', 'слушай', 'подскажи', 'плиз', 'пожалуйста',
    'пж', 'че', 'что', 'такое', 'работа', 'работу', 'рис', 'атлас', 'атласе',
    'если', 'есть', 'ли', 'бы', 'найди', 'нашел', 'нашла', 'скинь', 'пришли', 'посмотри', 'посмотреть', 'глянь', 'нарисуй',
    'да', 'ней', 'нем', 'нём', 'них', 'ей', 'ему', 'им', 'ними', 'все', 'всё', 'это',
    'этом', 'этой', 'эти', 'этих', 'этого', 'этим', 'быть', 'также', 'так', 'или', 'и',
    'из', 'за', 'над', 'под', 'при', 'от', 'до', 'без', 'через', 'со', 'собой'
}

GREETING_WORDS = {
    'привет', 'приветик', 'здравствуй', 'здравствуйте', 'хай', 'hello', 'hi',
    'добрый день', 'добрый вечер', 'доброе утро', 'пока', 'до свидания',
    'спасибо', 'благодарю', 'кто ты', 'как дела', 'что умеешь', 'помоги'
}

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = SCRIPT_DIR
LOG_DIR = os.path.join(PROJECT_ROOT, "logs")
if not os.path.exists(LOG_DIR):
    os.makedirs(LOG_DIR, exist_ok=True)

LOG_FILE = os.path.join(LOG_DIR, "rag_debug.log")

def log_msg(msg: str):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    formatted = f"{timestamp} - {msg}"
    print(formatted, flush=True)
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(formatted + "\n")
            f.flush()
    except Exception as e:
        print(f"Ошибка записи лога: {e}", file=sys.stderr)

def get_db_connection():
    dsn = os.getenv("DATABASE_URL")
    if dsn:
        try:
            return psycopg2.connect(dsn)
        except Exception:
            try:
                local_dsn = dsn.replace("@db:5432", "@localhost:5433").replace("@db:", "@localhost:5433:").replace("@db/", "@localhost:5433/")
                return psycopg2.connect(local_dsn)
            except Exception:
                pass

    return psycopg2.connect(
        host=os.getenv("PGHOST", "localhost"),
        port=os.getenv("PGPORT", "5433"),
        user=os.getenv("PGUSER", "postgres"),
        password=os.getenv("PGPASSWORD", "1010"),
        dbname=os.getenv("PGDATABASE", "medai_db")
    )

def fix_anatomical_typos(text: str) -> str:
    """Нормализация частых опечаток в ключевых словах"""
    text = re.sub(r'\bпока\b', 'покажи', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(д[иаео]+[афгрэл]+р[афгрэл]+гм\w*)\b', 'диафрагма', text, flags=re.IGNORECASE)
    return text

def is_greeting_or_smalltalk(text: str) -> bool:
    clean = re.sub(r'[^\w\s]', '', text.strip().lower())
    if not clean:
        return True
    words = clean.split()
    if all(w in GREETING_WORDS for w in words):
        return True
    return False

def get_russian_stem(text: str) -> str:
    words = re.split(r'\s+', text.strip().lower())
    words = [w for w in words if w]
    if not words:
        return ''

    stemmed = []
    suffix_pattern = re.compile(
        r'(ая|яя|ое|ее|ые|ие|ый|ий|ой|ей|ого|его|ому|ему|ым|им|ом|ем|ую|юю|ами|ями|ыми|ними|ах|ях|ов|ев|ей|ам|ям|а|я|о|е|ы|и|ь|у|ю)$',
        re.IGNORECASE
    )
    for word in words:
        if len(word) <= 3:
            stemmed.append(word)
        else:
            stemmed.append(suffix_pattern.sub('', word))
    return ' '.join(stemmed)

def clean_query_text(effective_query: str) -> str:
    cleaned = re.sub(r'[^\w\s]', ' ', effective_query.lower())
    words = [w for w in cleaned.split() if w]
    filtered = [w for w in words if w not in STOP_WORDS]
    if not filtered:
        return effective_query.strip()
    return ' '.join(filtered)

def is_latin_language_query(query: str) -> bool:
    q = query.lower()
    latin_triggers = [
        'латын', 'латинск', 'падеж', 'склонен', 'грамматик', 'предлог',
        'суффикс', 'приставк', 'терминоэлемент', 'морфем', 'словообразовани',
        'окончани', 'правило', 'грамматическ', 'согласован', 'чернявски',
        'на латыни', 'на латинском', 'как на латинском', 'как будет на латинском',
        'переведи на латынь', 'словарная форма'
    ]
    return any(trigger in q for trigger in latin_triggers)

def detect_intent(query: str) -> Dict[str, bool]:
    clean_query = query.lower().strip()

    if is_greeting_or_smalltalk(clean_query):
        return {"wants_image": False, "wants_text": True}

    words = [w for w in re.split(r'\s+', clean_query) if w]
    word_count = len(words)

    image_triggers = ['покажи', 'фото', 'картинка', 'картинку', 'фотка', 'фотку', 'фотке', 'рисунок', 'рисунки', 'схема', 'схему', 'иллюстрация', 'рис', 'изображение', 'нарисуй', 'скинь', 'пришли']
    text_triggers = [
        'объясни', 'расскажи', 'функция', 'функции', 'строение', 'патология', 'принцип', 'опиши', 'теория',
        'падеж', 'склонение', 'правило', 'перевод', 'предлог', 'суффикс', 'корень',
        'причины', 'этиология', 'симптом', 'лечение', 'разбор', 'где находится', 'как выглядит'
    ]

    wants_image = any(trig in clean_query for trig in image_triggers)
    wants_text = any(trig in clean_query for trig in text_triggers)

    if not wants_image and not wants_text:
        if word_count <= 3:
            wants_image = True
            wants_text = True
        else:
            wants_text = True

    return {"wants_image": wants_image, "wants_text": wants_text}

def has_contextual_pronouns(query: str) -> bool:
    q = f" {query.lower()} "
    pronouns = [
        ' их ', ' его ', ' её ', ' ей ', ' ему ', ' им ', ' ними ', ' они ',
        ' это ', ' этого ', ' этой ', ' этим ', ' этих ', ' данных ', ' данного ',
        ' у него ', ' у нее ', ' у них ', ' у мужчин', ' у женщин',
        ' картинку', ' картинка', ' фотке', ' фотка', ' фотку', ' фото', ' рисунок', ' схему', ' схеме', ' изображение'
    ]
    if any(p in q for p in pronouns):
        return True

    clean = clean_query_text(query)
    if len(clean) < 4:
        return True

    clean_words = [w for w in re.split(r'\s+', clean) if len(w) >= 4 and w not in STOP_WORDS]
    if not clean_words:
        return True

    return False

def extract_context_subject(db_messages: list, conn) -> Optional[str]:
    if not db_messages:
        return None

    count = len(db_messages)
    for i in range(count - 2, -1, -1):
        msg = db_messages[i]
        if msg.get('role') == 'user':
            parts = msg.get('parts', [])
            if isinstance(parts, str):
                try:
                    parts = json.loads(parts)
                except Exception:
                    parts = []

            prev_text = ""
            if isinstance(parts, list):
                for p in parts:
                    if isinstance(p, dict) and 'text' in p:
                        prev_text += p['text'] + " "

            prev_text = prev_text.strip()
            if not prev_text:
                continue

            clean_prev = clean_query_text(prev_text)

            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cur:
                    cur.execute("""
                        SELECT term_ru, term_lat 
                        FROM anatomy_terms 
                        WHERE term_ru ILIKE %s OR term_lat ILIKE %s
                        LIMIT 1
                    """, (f"%{clean_prev}%", f"%{clean_prev}%"))
                    row = cur.fetchone()
                    if row:
                        return row['term_ru']
            except Exception:
                try:
                    conn.rollback()
                except Exception:
                    pass

            if len(clean_prev) >= 3:
                return clean_prev

    return None

def get_yandex_embedding(text: str) -> Optional[List[float]]:
    if not YANDEX_API_KEY or not YANDEX_FOLDER_ID:
        log_msg("Предупреждение: YANDEX_CLOUD_API_KEY или YANDEX_CLOUD_FOLDER не заданы.")
        return None

    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/textEmbedding"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Api-Key {YANDEX_API_KEY}",
        "x-folder-id": YANDEX_FOLDER_ID
    }
    payload = {
        "modelUri": f"emb://{YANDEX_FOLDER_ID}/text-search-query/latest",
        "text": text[:2000]
    }

    try:
        with httpx.Client(timeout=10.0, verify=False) as client:
            resp = client.post(url, headers=headers, json=payload)
            if resp.status_code == 200:
                data = resp.json()
                return data.get("embedding")
            else:
                log_msg(f"Ошибка Yandex Embeddings HTTP {resp.status_code}: {resp.text}")
    except Exception as e:
        log_msg(f"Исключение при получении эмбеддинга Yandex: {e}")
    return None

def search_patient_dialogues(conn, query: str, limit: int = 2, query_vec: Optional[List[float]] = None) -> dict:
    """Выполняет векторный и текстовый RAG-поиск в таблице patient_dialogues (1 701 выверенный клинический диалог)"""
    result = {'context': '', 'count': 0}
    try:
        if query_vec is None:
            query_vec = get_yandex_embedding(query)

        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            if query_vec:
                vector_str = f"[{','.join(map(str, query_vec))}]"
                cur.execute("""
                    SELECT section_header, section_header_ru, section_text_ru, dialogue_text_ru,
                           (1.0 - (embedding <=> %s::vector)) AS vec_sim,
                           (
                               (1.0 - (embedding <=> %s::vector)) +
                               (CASE WHEN section_header IN ('GENHX', 'CC', 'ROS', 'EXAM', 'DIAGNOSIS', 'ASSESSMENT') THEN 0.10 ELSE 0.0 END)
                           ) AS rank_score
                    FROM patient_dialogues
                    WHERE embedding IS NOT NULL
                      AND (1.0 - (embedding <=> %s::vector)) >= 0.30
                    ORDER BY rank_score DESC
                    LIMIT %s
                """, (vector_str, vector_str, vector_str, limit))
            else:
                clean_q = clean_query_text(query)
                cur.execute("""
                    SELECT section_header, section_header_ru, section_text_ru, dialogue_text_ru, 0.0 AS vec_sim
                    FROM patient_dialogues
                    WHERE embedding IS NOT NULL
                      AND (section_header_ru ILIKE %s OR section_text_ru ILIKE %s OR dialogue_text_ru ILIKE %s)
                    ORDER BY (CASE WHEN section_header IN ('GENHX', 'CC', 'ROS', 'EXAM', 'DIAGNOSIS', 'ASSESSMENT') THEN 1 ELSE 2 END) ASC
                    LIMIT %s
                """, (f"%{clean_q}%", f"%{clean_q}%", f"%{clean_q}%", limit))

            rows = cur.fetchall()
            if rows:
                result['count'] = len(rows)
                ctx = "\n[КЛИНИЧЕСКИЕ ПРИМЕРЫ ИЗ БАЗЫ ПАЦИЕНТСКИХ ДИАЛОГОВ (patient_dialogues)]:\n"
                for i, r in enumerate(rows):
                    sim_val = round(float(r.get('vec_sim') or 0.0), 4)
                    header = r.get('section_header_ru', '')
                    sec_code = r.get('section_header', '')
                    summary = r.get('section_text_ru', '')
                    dialogue = r.get('dialogue_text_ru', '')
                    log_msg(f"  -> [Клинический диалог {i+1}] Категория [{sec_code}]: '{header}' | Векторное сходство: {sim_val}")
                    ctx += f"--- Клинический случай #{i+1} ({header}) ---\n"
                    if dialogue:
                        ctx += f"Пример общения с пациентом:\n{dialogue[:400]}\n"
                    if summary:
                        ctx += f"Врачебное заключение: {summary}\n"
                    ctx += "\n"
                result['context'] = ctx
    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        log_msg(f"Предупреждение: ошибка поиска в patient_dialogues: {e}")
    return result

def extract_medical_terms_via_gigachat(query: str) -> Optional[str]:
    """Быстрый микро-запрос к GigaChat-2 для очистки предложения и выделения канонического анатомического/медицинского термина."""
    if not query or len(query.strip()) < 4:
        return None

    try:
        t0 = datetime.now()
        with GigaChat(credentials=GIGACHAT_AUTH_KEY, scope=GIGACHAT_SCOPE, verify_ssl_certs=False, timeout=3.0) as giga:
            payload = Chat(
                model="GigaChat-2",
                messages=[
                    Messages(
                        role=MessagesRole.SYSTEM,
                        content=(
                            "Ты — медицинский лингвистический анализатор. Твоя задача — выделить ключевой "
                            "анатомический или медицинский термин из запроса пользователя и привести его к именительному падежу. "
                            "Ответь ТОЛЬКО выделенным термином, без лишних слов, кавычек и знаков препинания. Если вопрос по латыни, то ничего не меняй в нем и выведи полный текст запроса"
                        )
                    ),
                    Messages(role=MessagesRole.USER, content=query)
                ],
                temperature=0.01,
                max_tokens=40
            )
            resp = giga.chat(payload)
            extracted = (resp.choices[0].message.content or "").strip()
            extracted = re.sub(r'^[«"\'\`\s]+|[»"\'\`\s]+$', '', extracted)
            elapsed = (datetime.now() - t0).total_seconds()

            if extracted and len(extracted) >= 3 and not is_greeting_or_smalltalk(extracted):
                log_msg(f"[GigaChat-2 Term Extractor] Успешно выделен термин: '{extracted}' (время: {elapsed:.2f} сек)")
                return extracted
    except Exception as e:
        log_msg(f"Предупреждение: не удалось извлечь термин через GigaChat-2 ({e}). Используется локальный фильтр.")
    return None

def search_vector_knowledge_base(conn, user_query: str, mode: str = 'student', limit: int = 3, db_messages: list = []) -> dict:
    result = {
        'context': '',
        'image': None,
        'image_not_found': False,
        'wants_image': False,
        'wants_text': False,
        'is_latin_query': False
    }

    raw_query = user_query.strip()
    if not raw_query:
        log_msg("Запрос пустой.")
        return result

    trimmed_query = fix_anatomical_typos(raw_query)

    log_msg("=================== НОВЫЙ ЗАПРОС ПОЛЬЗОВАТЕЛЯ ===================")
    log_msg(f"Запрос: '{raw_query}' -> Исправлен: '{trimmed_query}' | Режим: '{mode}'")

    if is_greeting_or_smalltalk(trimmed_query):
        log_msg(f"Запрос '{trimmed_query}' распознан как приветствие. Поиск по БД пропущен.")
        result['wants_text'] = True
        return result

    effective_query = trimmed_query
    context_subject = None

    if has_contextual_pronouns(trimmed_query) and db_messages:
        context_subject = extract_context_subject(db_messages, conn)
        if context_subject:
            effective_query = f"{context_subject} {trimmed_query}"
            log_msg(f"Обнаружен контекстный запрос. Извлечен предмет из истории: '{context_subject}'")
            log_msg(f"Эффективный запрос для RAG: '{effective_query}'")

    is_latin = is_latin_language_query(effective_query) if mode == 'student' else False
    result['is_latin_query'] = is_latin

    intent = detect_intent(trimmed_query)
    wants_image = intent['wants_image'] if mode == 'student' else False
    wants_text = intent['wants_text']
    result['wants_image'] = wants_image
    result['wants_text'] = wants_text

    log_msg(f"Намерение -> Показать картинку: {'ДА' if wants_image else 'НЕТ'} | Текст: {'ДА' if wants_text else 'НЕТ'}")

    extracted_term = extract_medical_terms_via_gigachat(effective_query)
    if extracted_term:
        clean_search_query = clean_query_text(extracted_term)
        log_msg(f"Поисковый запрос RAG очищен через GigaChat-2: '{effective_query}' -> '{extracted_term}' (для поиска: '{clean_search_query}')")
    else:
        clean_search_query = clean_query_text(effective_query)

    stemmed_query = get_russian_stem(clean_search_query)

    # =========================================================================
    # РЕЖИМ ПАЦИЕНТА: БЕЗ ЛАТЫНИ И БЕЗ АНАТОМИЧЕСКИХ ИЛЛЮСТРАЦИЙ / КАРТИНОК
    # =========================================================================
    if mode == 'patient':
        log_msg("=== РЕЖИМ ПАЦИЕНТА: Поиск по словарям латыни (latin_vocab) и вывод фото из атласа анатомии (figure) ЗАБЛОКИРОВАНЫ ===")
        log_msg("Выполняется векторный поиск в базе клинических диалогов (patient_dialogues)...")

        # Получаем эмбеддинг эффективного запроса один раз для всех поисков
        query_vector = get_yandex_embedding(clean_search_query or effective_query)

        pat_res = search_patient_dialogues(conn, effective_query, limit=2, query_vec=query_vector)
        if pat_res['context']:
            result['context'] += pat_res['context']
            log_msg(f"Режим Пациента: успешно получено {pat_res['count']} релевантных диалогов из patient_dialogues.")
        else:
            log_msg("Режим Пациента: прямых совпадений в patient_dialogues не найдено.")

        # Дополнительный поиск только по текстовым материалам (book_text), без иллюстраций (figure)
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                if query_vector:
                    vector_str = f"[{','.join(map(str, query_vector))}]"
                    cur.execute("""
                        SELECT title, content, source_type, page_num,
                               (1.0 - (embedding <=> %s::vector)) AS vec_sim
                        FROM knowledge_base
                        WHERE source_type = 'book_text'
                          AND (1.0 - (embedding <=> %s::vector)) >= 0.32
                        ORDER BY vec_sim DESC
                        LIMIT %s
                    """, (vector_str, vector_str, limit))
                else:
                    cur.execute("""
                        SELECT title, content, source_type, page_num, 0.0 AS vec_sim
                        FROM knowledge_base
                        WHERE source_type = 'book_text'
                          AND (title ILIKE %s OR content ILIKE %s)
                        LIMIT %s
                    """, (f"%{clean_search_query}%", f"%{clean_search_query}%", limit))

                rows = cur.fetchall()
                if rows:
                    log_msg(f"Режим Пациента: получено {len(rows)} доп. текстовых записей из справочников (book_text).")
                    for i, r in enumerate(rows):
                        pnum = r.get('page_num') or 0
                        ctitle = r.get('title', '')
                        vec_sim = round(float(r.get('vec_sim', 0)), 4)
                        log_msg(f"  -> [Справочник {i+1}] Стр: {pnum} | Вектор: {vec_sim} | Заголовок: '{ctitle}'")
                        result['context'] += f"\n--- [СПРАВОЧНИК МЕДИЦИНЫ (Стр. {pnum})]: {ctitle} ---\n{r.get('content', '')}\n"
        except Exception as e:
            try: conn.rollback()
            except Exception: pass
            log_msg(f"Предупреждение: ошибка поиска в knowledge_base для Режима Пациента: {e}")

        result['image'] = None
        result['wants_image'] = False
        result['image_not_found'] = False
        return result

    # =========================================================================
    # РЕЖИМ СТУДЕНТА: АКАДЕМИЧЕСКИЙ ПОИСК С ЛАТЫНЬЮ И ТЕРМИНАМИ
    # =========================================================================
    log_msg(f"Область знания: {'УЧЕБНИК ЛАТЫНИ / ПЕРЕВОД' if is_latin else 'АНАТОМИЯ (учебник/атлас)'}")
    stem_words = [w for w in stemmed_query.split() if len(w) >= 3 and w not in STOP_WORDS]
    log_msg(f"Очищенный ключевой запрос для поиска: '{clean_search_query}' (стемминг: '{stemmed_query}')")

    matched_term_ru = context_subject
    matched_term_lat = None

    # Шаг 1: Поиск в anatomy_terms с гибким совпадением по термину
    if not matched_term_ru:
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("""
                    SELECT canonical_name, term_ru, term_lat,
                           GREATEST(
                                similarity(term_ru, %s),
                                similarity(COALESCE(term_lat, ''), %s),
                                similarity(term_ru, %s)
                           ) as sim
                    FROM anatomy_terms
                    WHERE (term_ru ILIKE %s OR term_lat ILIKE %s OR (LENGTH(term_ru) >= 3 AND %s ILIKE '%%' || term_ru || '%%'))
                       OR (similarity(term_ru, %s) >= 0.35 AND LENGTH(term_ru) >= 3)
                    ORDER BY sim DESC
                    LIMIT 1
                """, (
                    clean_search_query, clean_search_query, stemmed_query,
                    f"%{clean_search_query}%", f"%{clean_search_query}%", clean_search_query,
                    clean_search_query
                ))
                row = cur.fetchone()
                if row:
                    matched_term_ru = row['term_ru']
                    matched_term_lat = row['term_lat']
                    sim_val = round(float(row.get('sim', 0)), 4)
                    log_msg(f"Шаг 1: anatomy_terms распознали: '{clean_search_query}' -> RU: '{matched_term_ru}', LAT: '{matched_term_lat}' (Sim: {sim_val})")
                else:
                    log_msg("Шаг 1: В словаре anatomy_terms прямого совпадения фразы не найдено.")
        except Exception as e:
            try: conn.rollback()
            except Exception: pass
            log_msg(f"Предупреждение: ошибка обращения к anatomy_terms: {e}")

    vector_text = f"{matched_term_ru or clean_search_query} {matched_term_lat or ''}".strip()
    query_vector = get_yandex_embedding(vector_text)

    if not matched_term_ru and not is_latin:
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("""
                    SELECT title, content, figure_labels
                    FROM knowledge_base
                    WHERE source_type = 'latin_vocab'
                      AND (
                          title ILIKE %s OR content ILIKE %s
                       OR similarity(title, %s) >= 0.40
                      )
                    ORDER BY GREATEST(similarity(title, %s), similarity(content, %s)) DESC
                    LIMIT 1
                """, (f"%{clean_search_query}%", f"%{clean_search_query}%", clean_search_query, clean_search_query, clean_search_query))
                v_row = cur.fetchone()
                if v_row:
                    matched_term_ru = v_row['title']
                    matched_term_lat = v_row.get('figure_labels', '')
                    log_msg(f"Шаг 2: В latin_vocab гибридным поиском найден термин -> RU: '{matched_term_ru}', LAT: '{matched_term_lat}'")
                else:
                    log_msg("Шаг 2: В latin_vocab гибридным поиском ничего не найдено.")
        except Exception as e:
            try: conn.rollback()
            except Exception: pass
            log_msg(f"Предупреждение: ошибка гибридного поиска в latin_vocab: {e}")

    term_ru = matched_term_ru or clean_search_query
    term_lat = matched_term_lat or clean_search_query
    stem_ru_main = get_russian_stem(term_ru).split()[0] if term_ru else ""

    like_ru = f"%{term_ru}%"
    like_lat = f"%{term_lat}%"
    like_stem = f"%{stem_ru_main}%" if stem_ru_main else like_ru

    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            if query_vector:
                vector_str = f"[{','.join(map(str, query_vector))}]"
                sql = """
                    SELECT title, content, image_data, figure_info, figure_labels, source_type, page_num, vec_sim
                    FROM (
                        SELECT DISTINCT ON (title)
                            title, content, image_data, figure_info, figure_labels, source_type, page_num,
                            (1.0 - (embedding <=> %s::vector)) AS vec_sim,
                            (
                                ((1.0 - (embedding <=> %s::vector)) * 2.5)
                                + (CASE WHEN title ILIKE %s OR title ILIKE %s OR title ILIKE %s THEN 3.0 ELSE 0.0 END)
                                + (CASE WHEN figure_labels ILIKE %s OR figure_labels ILIKE %s THEN 2.5 ELSE 0.0 END)
                                + (CASE WHEN content ILIKE %s OR content ILIKE %s THEN 1.5 ELSE 0.0 END)
                            ) AS rank_score
                        FROM knowledge_base
                        ORDER BY title, rank_score DESC
                    ) dedup
                    ORDER BY (CASE WHEN source_type = 'book_text' THEN 1 WHEN source_type = 'figure' THEN 2 ELSE 3 END) ASC, rank_score DESC
                    LIMIT 10
                """
                cur.execute(sql, (vector_str, vector_str, like_ru, like_lat, like_stem, like_ru, like_lat, like_ru, like_lat))
            else:
                sql = """
                    SELECT title, content, image_data, figure_info, figure_labels, source_type, page_num, 0.0 AS vec_sim
                    FROM knowledge_base
                    WHERE title ILIKE %s OR title ILIKE %s OR title ILIKE %s OR figure_labels ILIKE %s OR content ILIKE %s
                    ORDER BY (CASE WHEN source_type = 'book_text' THEN 1 WHEN source_type = 'figure' THEN 2 ELSE 3 END) ASC
                    LIMIT 10
                """
                cur.execute(sql, (like_ru, like_lat, like_stem, like_ru, like_ru))

            raw_rows = cur.fetchall()

            if not raw_rows:
                all_words = list({w for w in re.split(r'\s+', f"{clean_search_query} {term_ru} {term_lat}") if len(w) >= 4 and w not in STOP_WORDS} | set(stem_words))
                all_words.sort(key=len, reverse=True)
                for word in all_words:
                    word_stem = get_russian_stem(word) or word
                    cur.execute("""
                        SELECT title, content, image_data, figure_info, figure_labels, source_type, page_num, 0.0 AS vec_sim
                        FROM knowledge_base
                        WHERE title ILIKE %s OR figure_labels ILIKE %s OR content ILIKE %s OR title ILIKE %s
                        ORDER BY (CASE WHEN source_type = 'figure' THEN 2 ELSE 1 END) ASC
                        LIMIT 10
                    """, (f"%{word}%", f"%{word}%", f"%{word}%", f"%{word_stem}%"))
                    raw_rows = cur.fetchall()
                    if raw_rows:
                        log_msg(f"Шаг 3: text-fallback по слову '{word}' ({word_stem}) → {len(raw_rows)} записей")
                        break

        rows = raw_rows[:limit] if raw_rows else []

        if rows:
            log_msg(f"Шаг 3: Сбалансировано и получено материалов: {len(rows)}")
            context_text = "\n\n[СПРАВОЧНЫЕ МАТЕРИАЛЫ ИЗ БАЗЫ ЗНАНИЙ]:\n"

            for i, r in enumerate(rows):
                source_type = r.get('source_type') or 'general'
                page_num = r.get('page_num') or 0
                clean_title = re.sub(r'^Страница\s+\d+\s+—\s+Рис\.\s*\d+:\s*', '', r.get('title', ''), flags=re.IGNORECASE)
                content = r.get('content', '')
                vec_sim = round(float(r.get('vec_sim', 0)), 4)

                log_msg(f"  -> [Запись {i+1}] Категория: '{source_type}' | Стр: {page_num} | Вектор: {vec_sim} | Заголовок: '{clean_title}'")
                context_text += f"--- [{source_type.upper()} (Стр. {page_num})]: {clean_title} ---\n{content}\n\n"

                if wants_image and result['image'] is None and r.get('image_data') and r.get('source_type') == 'figure':
                    result['image'] = r['image_data']
                    log_msg(f"     Выбрано РЕЛЕВАНТНОЕ изображение: {clean_title}")

            result['context'] = context_text
        else:
            log_msg("Шаг 3: Поиск в учебнике не дал результатов.")

        if wants_image and result['image'] is None:
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cur:
                    cur.execute("""
                        SELECT title, image_data, page_num
                        FROM knowledge_base
                        WHERE source_type = 'figure'
                          AND image_data IS NOT NULL
                          AND (title ILIKE %s OR figure_labels ILIKE %s OR content ILIKE %s)
                        ORDER BY page_num ASC
                        LIMIT 1
                    """, (like_ru, like_ru, like_ru))
                    fig_row = cur.fetchone()
                    if fig_row and fig_row.get('image_data'):
                        result['image'] = fig_row['image_data']
                        log_msg(f"Выбрано РЕЛЕВАНТНОЕ изображение из доп. запроса рисунков: {fig_row['title']}")
            except Exception as e_fig:
                log_msg(f"Предупреждение: ошибка доп. запроса рисунка: {e_fig}")

        if wants_image and result['image'] is None:
            result['image_not_found'] = True
            log_msg("Флаг: Попросили картинку, но релевантное изображение в БД НЕ найдено.")

    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        log_msg(f"Ошибка гибридного запроса SQL: {e}")

    return result

def extract_and_match_interactive_terms(ai_response: str, conn) -> Dict[str, Any]:
    """Извлекает латинские термины из ответа нейросети и ищет для них иллюстрации в knowledge_base"""
    if not ai_response:
        return {}

    latin_matches = re.findall(r'\(\*?([a-zA-Z\s\-]{3,60})\*?\)', ai_response)
    standalone_latin = re.findall(r'\b([a-zA-Z]{3,25}(?:\s+[a-zA-Z]{3,25}){1,3})\b', ai_response)
    
    candidates = list(dict.fromkeys(latin_matches + standalone_latin))
    IGNORE_WORDS = {'giga', 'chat', 'http', 'https', 'null', 'none', 'true', 'false', 'text', 'mode', 'terminologia', 'anatomica', 'latin'}
    GENERIC_TERMS = {'pars', 'partes', 'ossa', 'arteria', 'arteriae', 'veina', 'venae', 'nervus', 'nervi', 'musculus', 'musculi', 'ligamentum', 'ligamenta', 'tractus', 'fossa', 'sulcus', 'canalis', 'foramen', 'processus', 'tuber', 'bursa', 'articulatio', 'regio', 'facies', 'corpus', 'caput', 'margo', 'spina', 'arcus', 'lamina', 'incisura', 'fissura', 'hiatus', 'sinus', 'plexus', 'ramus', 'rami', 'nodi', 'glandula', 'truncus', 'nucleus', 'nuclei', 'lobus', 'gyrus'}
    
    interactive_terms = {}
    
    for cand in candidates:
        term_clean = cand.strip()
        if len(term_clean) < 3 or term_clean.lower() in IGNORE_WORDS:
            continue
            
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                # Строгое совпадение по границам всего термина (без ложных совпадений по отдельным общим словам)
                regex_pattern = r'\y' + r'\s+'.join(re.escape(w) for w in re.split(r'\s+', term_clean)) + r'\y'
                cur.execute("""
                    SELECT title, page_num, image_data
                    FROM knowledge_base
                    WHERE source_type = 'figure'
                      AND image_data IS NOT NULL
                      AND (
                          title ~* %s OR figure_labels ~* %s OR content ~* %s
                      )
                    ORDER BY page_num ASC
                    LIMIT 1
                """, (regex_pattern, regex_pattern, regex_pattern))
                row = cur.fetchone()

                if row and row.get('image_data'):
                    clean_title = re.sub(r'^Страница\s+\d+\s+—\s+Рис\.\s*\d+:\s*', '', row['title'], flags=re.IGNORECASE)
                    interactive_terms[term_clean] = {
                        "image_url": row['image_data'],
                        "title": clean_title,
                        "page_num": row.get('page_num', 0)
                    }
        except Exception as e:
            # Здесь можно логировать ошибку БД, если нужно
            print(f"Ошибка при поиске термина {term_clean}: {e}")
            continue

    # Исключаем подстроки (например, 'kyphosis sacralis' убираем, если есть 'kyphosis sacralis et coccygea')
    filtered_terms = {}
    sorted_matched_keys = sorted(interactive_terms.keys(), key=lambda k: len(k), reverse=True)
    for key in sorted_matched_keys:
        is_substring = False
        for existing in filtered_terms:
            if key.lower() in existing.lower():
                is_substring = True
                break
        if not is_substring:
            filtered_terms[key] = interactive_terms[key]

    return filtered_terms


def get_gigachat_token() -> str:
    url = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
    headers = {
        "Content-Type": "application/x-www-form-urlencoded",
        "Accept": "application/json",
        "RqUID": str(uuid.uuid4()),
        "Authorization": f"Basic {GIGACHAT_AUTH_KEY}"
    }
    data = {"scope": GIGACHAT_SCOPE}

    with httpx.Client(timeout=15.0, verify=False) as client:
        resp = client.post(url, headers=headers, data=data)
        resp.raise_for_status()
        return resp.json()["access_token"]

def extract_text_from_file(file_bytes: bytes, filename: str) -> Optional[str]:
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.txt', '.csv', '.md', '.json', '.log', '.html', '.xml']:
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_bytes.decode('cp1251')
            except Exception:
                return None
    elif ext == '.pdf':
        try:
            import io
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
            return text.strip() if text.strip() else None
        except Exception as e:
            log_msg(f"  [PDF] Не удалось прочитать '{filename}': {e}")
            return None
    elif ext in ['.docx', '.doc']:
        # 1. Попытка вычитывания через python-docx
        try:
            import io
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_txt:
                        paragraphs.append(row_txt)
            text = "\n".join(paragraphs).strip()
            if text:
                return text
        except Exception as e:
            log_msg(f"  [DOCX] docx библиотека не вычитала '{filename}': {e}")

        # 2. Гарантированный XML Fallback через zipfile (работает всегда на стандартных .docx)
        try:
            import io, zipfile, xml.etree.ElementTree as ET
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    tree = ET.fromstring(xml_content)
                    texts = []
                    for node in tree.iter():
                        if node.tag.endswith('}t') and node.text:
                            texts.append(node.text)
                        elif node.tag.endswith('}p'):
                            texts.append("\n")
                    full_text = "".join(texts).strip()
                    if full_text:
                        return full_text
        except Exception as e:
            log_msg(f"  [DOCX XML Fallback] Ошибка вычитывания word/document.xml: {e}")
            return None
    return None

def upload_file_to_gigachat(token: str, file_bytes: bytes, filename: str) -> Optional[str]:
    url = "https://api.giga.chat/v1/files"
    headers = {"Authorization": f"Bearer {token}"}
    files = {"file": (filename, file_bytes)}
    data = {"purpose": "general"}
    try:
        with httpx.Client(timeout=30.0, verify=False) as client:
            resp = client.post(url, headers=headers, data=data, files=files)
            if resp.status_code == 200:
                file_id = resp.json().get("id")
                log_msg(f"  -> Файл '{filename}' успешного загружен в GigaChat Storage. ID: {file_id}")
                return file_id
            else:
                log_msg(f"  -> Ошибка загрузки файла '{filename}': HTTP {resp.status_code} - {resp.text}")
    except Exception as e:
        log_msg(f"  -> Исключение при загрузке файла '{filename}': {e}")
    return None

@app.post("/api/chat")
async def process_chat(
    request: Request,
    query: str = Form(""),
    mode: str = Form("student"),
    university: str = Form(""),
    chat_id: str = Form(""),
    user_login: str = Form(""),
    files: List[UploadFile] = File(default=[])
):
    try:
        # Извлекаем ВСЕ загруженные файлы из формы и параметров
        form = await request.form()
        if not query or not query.strip():
            query = str(form.get("query") or form.get("user_query") or "").strip()

        all_files: List[Any] = []
        for key, value in form.multi_items():
            fn = getattr(value, "filename", None)
            if fn and value not in all_files:
                all_files.append(value)
        for f in files:
            fn = getattr(f, "filename", None)
            if fn and f not in all_files:
                all_files.append(f)

        log_msg("=================== НОВЫЙ ВХОДЯЩИЙ ЗАПРОС ===================")
        log_msg(f"Пользователь: '{user_login}' | Чат: '{chat_id}' | Режим: '{mode}' | ВУЗ: '{university}'")
        log_msg(f"Текст запроса: '{query}' | Прикреплено файлов: {len(all_files)}")

        conn = get_db_connection()

        is_guest = user_login.startswith("guest_") or user_login == "guest" or not user_login

        if not is_guest:
            if not chat_id:
                chat_id = f"chat_{uuid.uuid4().hex[:12]}"
                title = (query[:30] + '...') if query else 'Анализ документов...'
                with conn.cursor() as cur:
                    cur.execute(
                        "INSERT INTO chats (chat_id, user_login, mode, university, title, updated_at) VALUES (%s, %s, %s, %s, %s, NOW())",
                        (chat_id, user_login, mode, university, title)
                    )
                conn.commit()
            else:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE chats SET updated_at = NOW() WHERE chat_id = %s AND user_login = %s",
                        (chat_id, user_login)
                    )
                conn.commit()
        elif not chat_id:
            chat_id = f"chat_{uuid.uuid4().hex[:12]}"

        uploaded_parts = []
        uploaded_file_ids = []

        if all_files:
            try:
                giga_token = get_gigachat_token()
                for file in all_files:
                    file_bytes = await file.read()
                    if file_bytes:
                        log_msg(f"[Файл {file.filename}] Получен байтовый поток: {len(file_bytes)} байт, Content-Type: {file.content_type}")
                        file_id = upload_file_to_gigachat(giga_token, file_bytes, file.filename)
                        if file_id:
                            uploaded_file_ids.append(file_id)

                        extracted_text = extract_text_from_file(file_bytes, file.filename)
                        content_type = file.content_type or "application/octet-stream"
                        ext = os.path.splitext(file.filename)[1].lower()

                        if extracted_text:
                            log_msg(f"[Файл {file.filename}] Извлечен текстовый контент ({len(extracted_text)} символов)")
                            uploaded_parts.append({
                                "text": f"\n[СОДЕРЖИМОЕ ПРИКРЕПЛЕННОГО ДОКУМЕНТА '{file.filename}']:\n{extracted_text}\n[КОНЕЦ ДОКУМЕНТА]\n",
                                "file_name": file.filename
                            })
                        elif content_type.startswith('image/') or ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                            b64 = base64.b64encode(file_bytes).decode('utf-8')
                            log_msg(f"[Файл {file.filename}] Загружен как Base64 изображение (длина b64: {len(b64)})")
                            uploaded_parts.append({
                                "file_name": file.filename,
                                "inline_data": {
                                    "mime_type": content_type if content_type.startswith('image/') else 'image/jpeg',
                                    "data": b64
                                }
                            })
                        else:
                            log_msg(f"[Файл {file.filename}] Прикреплен как бинарный файл без прямого вычитывания текста")
                            uploaded_parts.append({
                                "text": f"\n[Прикреплен документ: {file.filename}]\n",
                                "file_name": file.filename
                            })
            except Exception as e:
                log_msg(f"Ошибка обработки прикрепленных файлов: {e}")

        current_parts = list(uploaded_parts)
        if query:
            current_parts.append({"text": query})

        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO messages (chat_id, role, parts) VALUES (%s, 'user', %s)",
                (chat_id, json.dumps(current_parts, ensure_ascii=False))
            )
        conn.commit()

        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT role, parts FROM messages WHERE chat_id = %s ORDER BY id ASC", (chat_id,))
            db_messages = cur.fetchall()

        rag_result = search_vector_knowledge_base(conn, query, mode=mode, limit=3, db_messages=db_messages)

        if mode == 'student':
            system_instruction = (
                "Ты — MedAI, виртуальный ИИ-наставник и академический эксперт по всем областям медицины. Твой стиль: предельная точность, "
                "сухой академический язык, точность в терминологии и отсутствие пустой «воды»."
                + (f" Студент из {university}." if university else "") + "\n"
            )
        else:
            system_instruction = (
                "Ты — практикующий врач-диагност и эксперт по медицинской терминологии. "
                "Будь точен, давай понятные, краткие рекомендации. Собирай информацию, задавая "
                "уточняющие вопросы, затем на основе нескольких сообщений дай приблизительный диагноз. Не используй латынь.\n"
            )

        system_instruction += "\nСТРОГИЕ ПРАВИЛА И СТИЛЬ ОТВЕТА:\n"
        system_instruction += (
            "1. ИСКЛЮЧИ «ВОДУ» И КЛИШЕ: КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНЫ вводные штампы: "
            "'играет важную роль', 'имеет большое значение', 'в клинической практике'. Отвечай сразу по существу.\n"
        )
        if mode == 'student':
            system_instruction += (
                "2. ОБЯЗАТЕЛЬНАЯ ЛАТЫНЬ (Terminologia Anatomica): При описании любых анатомических структур, костей, "
                "суставов, органов, мышц, сосудов и нервов ВСЕГДА дублируй их точные официальные латинские названия "
                "в скобках (например: череп (*cranium*), лобная кость (*os frontale*), большеберцовая кость (*tibia*)).\n"
            )
        else:
            system_instruction += (
                "2. СТРОГИЙ ЗАПРЕТ НА ЛАТЫНЬ: В режиме Пациента КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО использовать любые латинские термины "
                "и слова на латинице. Отвечай исключительно простым, понятным и сочувствующим языком на русском.\n"
            )
        system_instruction += (
            "3. СТРОГИЙ ЗАПРЕТ НА ОТГОВОРКИ: КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО писать фразы вроде "
            "'обратитесь к специализированным атласам', 'у меня нет возможности показать изображение', 'я лишь текстовая модель'. "
            "Давай полный академический ответ сам!\n"
        )
        system_instruction += (
            "4. ОФОРМЛЕНИЕ ТАБЛИЦ: Если ответ содержит данные из таблиц (категория 'latin_table' или "
            "таблицы терминоэлементов/склонений), ВСЕГДА оформляй их в виде аккуратной Markdown-таблицы с разделителями '|'.\n"
        )

        if rag_result["is_latin_query"]:
            system_instruction += (
                "5. ТРЕБОВАНИЯ К ЛАТИНСКИМ ПЕРЕВОДАМ: При явном запросе перевода на латынь давай полную словарную форму "
                "(*Nominativus*, *Genitivus*, род, склонение) и 2-3 примера из *Terminologia Anatomica*.\n"
            )
        else:
            system_instruction += (
                "5. АНАТОМИЧЕСКИЙ ЗАПРОС: Отвечай строго по строению и клинической анатомии. НЕ выводи словарные "
                "грамматические лекции (склонения), если пользователя интересует анатомический объект или его патология.\n"
            )

        if not rag_result["wants_text"]:
            system_instruction += (
                "6. КРАТКОСТЬ: Пользователь попросил только показать объект. Ответь предельно кратко "
                "(1–2 предложения с точным определением и латинским термином), без длинных лекций и разделов!\n"
            )

        system_instruction += (
            "7. ЗАПРЕТ НА ОПИСАНИЕ НОМЕРОВ РИСУНКОВ И СТРАНИЦ АТЛАСА: КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО писать отчёты "
            "по номерам иллюстраций или перечислять подписи к рисункам («На рисунке 108 показано...», «Полусхематичный вид...», «Рис. 1...»). "
            "Давай только прямой содержательный академический ответ по анатомическому строению, отделам, костям, каналам и функциям.\n"
        )

        if rag_result["context"]:
            system_instruction += (
                "\n\nСПРАВОЧНЫЕ ДАННЫЕ ИЗ БАЗЫ ЗНАНИЙ (Используй их как основу для ответа, извлеки суть):\n"
                + rag_result["context"]
            )
            if rag_result["image"]:
                system_instruction += "\n\nПРИМЕЧАНИЕ: Наглядное изображение выведено на экран пользователя. Дай точно его краткое описание."
        else:
            system_instruction += (
                "\n\n7. ГИБКИЙ ОТВЕТ: Если пользователь ведет общение (приветствия, вежливые обращения), "
                "отвечай вежливо и естественно на русском языке, приглашая задать анатомический вопрос. "
                "Если прикреплен файл без текста, дай его подробный анализ."
            )

        gigachat_messages = [
            Messages(role=MessagesRole.SYSTEM, content=system_instruction)
        ]

        total_msgs = len(db_messages)
        for idx, msg in enumerate(db_messages):
            role_str = msg.get('role')
            role = MessagesRole.ASSISTANT if role_str == 'model' else MessagesRole.USER
            parts = msg.get('parts', [])
            if isinstance(parts, str):
                try:
                    parts = json.loads(parts)
                except Exception:
                    parts = []

            text_acc = ""
            for p in parts:
                if isinstance(p, dict) and 'text' in p:
                    text_acc += p['text'] + "\n"

            text_acc = text_acc.strip()
            if not text_acc:
                text_acc = "Проанализируй прикрепленный документ." if uploaded_file_ids else "Запрос пользователя."

            if idx == total_msgs - 1 and uploaded_file_ids and role == MessagesRole.USER:
                gigachat_messages.append(
                    Messages(role=role, content=text_acc, attachments=uploaded_file_ids)
                )
            else:
                gigachat_messages.append(Messages(role=role, content=text_acc))

        log_msg(f"Отправка запроса в GigaChat SDK (модель: {GIGACHAT_MODEL}, scope: {GIGACHAT_SCOPE}, сообщений: {len(gigachat_messages)}, file_ids: {uploaded_file_ids})...")

        ai_response = ""
        max_retries = 2
        for attempt in range(1, max_retries + 1):
            try:
                with GigaChat(credentials=GIGACHAT_AUTH_KEY, scope=GIGACHAT_SCOPE, verify_ssl_certs=False, timeout=60.0) as giga:
                    payload = Chat(
                        model=GIGACHAT_MODEL,
                        messages=gigachat_messages,
                        temperature=0.4,
                        max_tokens=3000
                    )
                    response = giga.chat(payload)
                    ai_response = response.choices[0].message.content or ""
                    if ai_response:
                        break
            except Exception as giga_err:
                log_msg(f"Попытка {attempt}/{max_retries} запроса к GigaChat завершилась ошибкой: {giga_err}")
                if attempt == max_retries:
                    raise giga_err
                import time
                time.sleep(1.0)

        ai_response = re.sub(r'<(think|thought)>.*?</\1>', '', ai_response, flags=re.DOTALL).strip()

        if ai_response:
            log_msg(f"========= ОТВЕТ GIGACHAT (Длина: {len(ai_response)} симв.) =========")
            log_msg(ai_response)
            log_msg("================================================================")

            if mode == 'student':
                interactive_terms = extract_and_match_interactive_terms(ai_response, conn)
                if interactive_terms:
                    log_msg(f"Найдено {len(interactive_terms)} подсвечиваемых латинских терминов с иллюстрациями: {list(interactive_terms.keys())}")
            else:
                interactive_terms = {}
                rag_result["wants_image"] = False
                rag_result["image"] = None

            model_parts: List[Dict[str, Any]] = [{"text": ai_response}]
            if rag_result.get("wants_image") and rag_result.get("image"):
                raw_b64 = re.sub(r'^data:image/\w+;base64,', '', rag_result["image"])
                model_parts.append({
                    "inline_data": {"mime_type": "image/jpeg", "data": raw_b64},
                    "file_name": "illustration.jpg"
                })
            if interactive_terms:
                model_parts.append({"interactive_terms": interactive_terms})

            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO messages (chat_id, role, parts) VALUES (%s, 'model', %s)",
                    (chat_id, json.dumps(model_parts, ensure_ascii=False))
                )
            conn.commit()

            response_data = {"response": ai_response, "chat_id": chat_id}
            if rag_result.get("wants_image") and rag_result.get("image"):
                response_data["image_url"] = rag_result["image"]
            if interactive_terms:
                response_data["interactive_terms"] = interactive_terms

            return JSONResponse(content=response_data)
        else:
            log_msg("ОШИБКА: GigaChat вернул пустой ответ.")
            return JSONResponse(content={"response": "Ошибка: Нейросеть вернула пустой ответ.", "chat_id": chat_id})

    except Exception as e:
        log_msg(f"ОШИБКА ОБРАБОТКИ ЧАТА: {e}")
        return JSONResponse(content={"response": f"Ошибка обработки запроса ИИ: {str(e)}", "chat_id": chat_id}, status_code=500)

    finally:
        if 'conn' in locals() and conn:
            conn.close()